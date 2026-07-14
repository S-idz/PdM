"""
Converts docs/Chapter_8_Report.md to docs/Chapter_8_Report.docx, preserving
headings, tables, images, bold/italic inline text, bullet/numbered lists,
and code blocks. Purpose-built for this file's markdown shape (no general
CommonMark parser) -- run from repo root:

    python scripts/md_to_docx_chapter8.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "Chapter_8_Report.md"
DOCS_DIR = ROOT / "docs"
OUT = ROOT / "docs" / "Chapter_8_Report.docx"

INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def add_inline_runs(paragraph, text):
    """Handle **bold**, `code`, *italic* within one line of text."""
    # tokenize by splitting on bold/code/italic markers, preserving order
    tokens = []
    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*[^*]+?\*(?!\*))")
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append(("text", text[pos:m.start()]))
        tok = m.group(0)
        if tok.startswith("**"):
            tokens.append(("bold", tok[2:-2]))
        elif tok.startswith("`"):
            tokens.append(("code", tok[1:-1]))
        else:
            tokens.append(("italic", tok[1:-1]))
        pos = m.end()
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    if not tokens:
        tokens = [("text", text)]
    for kind, val in tokens:
        run = paragraph.add_run(val)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x2A, 0x2A)


def add_table(doc, rows):
    header, *body = rows
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        add_inline_runs(p, h)
        for run in p.runs:
            run.bold = True
    for r in body:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            if i < ncols:
                cells[i].text = ""
                add_inline_runs(cells[i].paragraphs[0], val)
    doc.add_paragraph()


def parse_table_block(lines, start):
    i = start
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s:\-|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # code fence
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                for cl in code_buf:
                    run = p.add_run(cl + "\n")
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # image
        img_m = IMG_RE.match(stripped)
        if img_m:
            alt, relpath = img_m.groups()
            img_path = (DOCS_DIR / relpath).resolve()
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.0))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[missing image: {relpath}]")
                run.italic = True
            i += 1
            continue

        # table
        if stripped.startswith("|"):
            rows, new_i = parse_table_block(lines, i)
            if rows:
                add_table(doc, rows)
            i = new_i
            continue

        # numbered list
        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m_num.group(2))
            i += 1
            continue

        # bullet list
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        # bold standalone line (e.g. **Figure X: ...**)
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
